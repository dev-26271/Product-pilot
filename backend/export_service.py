import os
import json
from datetime import datetime, timezone
from typing import Dict, Any, List, Tuple

# python-docx imports
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.pdfgen import canvas


class NumberedCanvas(canvas.Canvas):
    """Custom canvas to compute total page count and draw running headers/footers."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            # Suppress running header/footer on cover page
            return
        
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#7F8C8D"))
        
        # Header
        self.drawString(54, 755, "ProductPilot Workspace Export")
        self.setStrokeColor(colors.HexColor("#BDC3C7"))
        self.setLineWidth(0.5)
        self.line(54, 748, 558, 748)
        
        # Footer
        self.line(54, 45, 558, 45)
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 32, page_text)
        self.drawString(54, 32, "CONFIDENTIAL — FOR INTERNAL USE ONLY")
        self.restoreState()


class ExportService:
    """Handles serialization and formatting of workspace deliverables to multiple formats."""

    @staticmethod
    def get_export_filepath(project_name: str, scope: str, fmt: str) -> str:
        """Returns the absolute path to save the export file."""
        safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_")
        safe_scope = scope.lower().replace(" ", "_")
        
        # Exports directory in the workspace
        exports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "exports", safe_project_name))
        os.makedirs(exports_dir, exist_ok=True)
        
        filename = f"{safe_scope}.{fmt}"
        return os.path.join(exports_dir, filename)

    @classmethod
    def export(cls, workspace: Dict[str, Any], scope: str, fmt: str) -> Dict[str, Any]:
        """Runs the export pipeline for the target scope and format."""
        project_name = workspace.get("name", "ProductPilot_Project")
        filepath = cls.get_export_filepath(project_name, scope, fmt)
        
        # Resolve target data based on scope
        data_to_export = cls._resolve_scope_data(workspace, scope)
        if not data_to_export:
            return {
                "status": "failure",
                "error": f"No content generated yet for scope '{scope}'. Please generate it first."
            }
            
        try:
            if fmt == "json":
                cls._export_json(data_to_export, filepath)
            elif fmt == "md":
                cls._export_markdown(data_to_export, scope, workspace, filepath)
            elif fmt == "docx":
                cls._export_docx(data_to_export, scope, workspace, filepath)
            elif fmt == "pdf":
                cls._export_pdf(data_to_export, scope, workspace, filepath)
            else:
                return {
                    "status": "failure",
                    "error": f"Unsupported export format '{fmt}'"
                }
                
            return {
                "status": "success",
                "location": filepath,
                "filename": os.path.basename(filepath),
                "scope": scope,
                "format": fmt
            }
        except Exception as e:
            return {
                "status": "failure",
                "error": f"Export failed: {str(e)}"
            }

    @classmethod
    def _resolve_scope_data(cls, workspace: Dict[str, Any], scope: str) -> Any:
        """Resolves target dictionary, list, or context for the requested scope."""
        deliverables = workspace.get("deliverables", {})
        
        if scope == "Entire Workspace":
            # Only export if at least PRD is generated
            if "Product Requirements Document (PRD)" not in deliverables:
                return None
            return workspace
            
        elif scope == "PRD":
            return workspace.get("prd") or deliverables.get("Product Requirements Document (PRD)")
            
        elif scope == "BRD":
            return deliverables.get("Business Requirements Document (BRD)")
            
        elif scope == "SRS":
            return deliverables.get("Software Requirements Specification (SRS)")
            
        elif scope == "User Stories":
            return deliverables.get("User Stories")
            
        elif scope == "Roadmap":
            return deliverables.get("Product Roadmap")
            
        elif scope == "Jira Tasks":
            return deliverables.get("Jira Tasks")
            
        elif scope == "Sprint Backlog":
            return deliverables.get("Sprint Backlog")
            
        elif scope == "Executive Summary":
            # Extract executive summary from PRD if available
            prd = workspace.get("prd") or deliverables.get("Product Requirements Document (PRD)", {})
            content = prd.get("content", prd)
            if isinstance(content, dict):
                for k, v in content.items():
                    if "executive summary" in k.lower():
                        return {k: v}
            return {"Executive Summary": workspace.get("idea", "No summary generated.")}
            
        return None

    # ── Exporters ─────────────────────────────────────────────────────────────

    @classmethod
    def _export_json(cls, data: Any, filepath: str) -> None:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, default=str)

    @classmethod
    def _export_markdown(cls, data: Any, scope: str, workspace: Dict[str, Any], filepath: str) -> None:
        lines = []
        lines.append(f"# {workspace.get('name', 'ProductPilot Project')} - {scope}")
        lines.append(f"**Version:** {workspace.get('metadata', {}).get('version', '1.0.0')}")
        lines.append(f"**Generated:** {datetime.now(timezone.utc).isoformat()}")
        lines.append("\n---\n")
        
        lines.extend(cls._convert_to_markdown_lines(data, scope))
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    @classmethod
    def _export_docx(cls, data: Any, scope: str, workspace: Dict[str, Any], filepath: str) -> None:
        doc = DocxDocument()
        
        # Style variables
        title_color = RGBColor(30, 58, 138)  # Navy Blue
        text_color = RGBColor(51, 65, 85)   # Charcoal
        
        # 1. Cover Page
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(f"\n\n\n{workspace.get('name', 'ProductPilot Project')}")
        run_title.font.size = Pt(28)
        run_title.font.bold = True
        run_title.font.color.rgb = title_color
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(f"Workspace Deliverable Export — {scope}\n\n\n")
        run_sub.font.size = Pt(14)
        run_sub.font.italic = True
        run_sub.font.color.rgb = text_color
        
        # Logo Box Placeholder
        table_logo = doc.add_table(rows=1, cols=1)
        table_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table_logo.cell(0, 0)
        cell.width = Inches(3.0)
        p_logo = cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run("[ COMPANY LOGO PLACEHOLDER ]")
        run_logo.font.size = Pt(10)
        run_logo.font.bold = True
        run_logo.font.color.rgb = RGBColor(148, 163, 184)
        
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = (
            f"\n\n\n\nExport Metadata\n"
            f"Version: {workspace.get('metadata', {}).get('version', '1.0.0')}\n"
            f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}\n"
            f"Status: Active Deliverable"
        )
        p_meta.add_run(meta_text).font.color.rgb = text_color
        
        doc.add_page_break()
        
        # 2. Table of Contents Placeholder
        h_toc = doc.add_paragraph()
        h_toc.add_run("Table of Contents").font.size = Pt(16)
        h_toc.runs[0].font.bold = True
        h_toc.runs[0].font.color.rgb = title_color
        
        doc.add_paragraph("1. Project Overview\n2. Deliverable Details\n3. Quality Audit Validation")
        doc.add_page_break()
        
        # 3. Content Insertion
        cls._insert_docx_content(doc, data, scope)
        
        doc.save(filepath)

    @classmethod
    def _export_pdf(cls, data: Any, scope: str, workspace: Dict[str, Any], filepath: str) -> None:
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        style_title = ParagraphStyle(
            name='CoverTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=32,
            leading=38,
            textColor=colors.HexColor("#1E3A8A"),
            alignment=1, # Center
            spaceAfter=15
        )
        
        style_subtitle = ParagraphStyle(
            name='CoverSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#475569"),
            alignment=1,
            spaceAfter=40
        )
        
        style_heading = ParagraphStyle(
            name='SectionHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1E3A8A"),
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        )
        
        style_body = ParagraphStyle(
            name='BodyTextCustom',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#334155"),
            spaceAfter=10
        )
        
        story = []
        
        # 1. Cover Page
        story.append(Spacer(1, 100))
        story.append(Paragraph(workspace.get('name', 'ProductPilot Project'), style_title))
        story.append(Paragraph(f"Workspace Deliverable Export — {scope}", style_subtitle))
        story.append(Spacer(1, 40))
        
        # Logo Box Table
        logo_data = [["[ COMPANY LOGO PLACEHOLDER ]"]]
        t_logo = Table(logo_data, colWidths=[200], rowHeights=[60])
        t_logo.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#F1F5F9")),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor("#94A3B8")),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('INNERGRID', (0,0), (-1,-1), 1, colors.HexColor("#CBD5E1")),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#CBD5E1")),
        ]))
        story.append(t_logo)
        story.append(Spacer(1, 100))
        
        # Metadata Block
        meta_html = (
            f"<para align=center><b>Export Details</b><br/>"
            f"Version: {workspace.get('metadata', {}).get('version', '1.0.0')}<br/>"
            f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}<br/>"
            f"Status: Active Product Deliverable</para>"
        )
        story.append(Paragraph(meta_html, style_body))
        story.append(PageBreak())
        
        # 2. Table of Contents
        story.append(Paragraph("Table of Contents", style_heading))
        story.append(Spacer(1, 10))
        story.append(Paragraph("1. Project Details Overview", style_body))
        story.append(Paragraph("2. Scope Specifications", style_body))
        story.append(Paragraph("3. Functional Breakdown", style_body))
        story.append(PageBreak())
        
        # 3. Add Content Flowables
        cls._insert_pdf_content(story, data, scope, style_heading, style_body)
        
        doc.build(story, canvasmaker=NumberedCanvas)

    # ── Content converters ───────────────────────────────────────────────────

    @classmethod
    def _convert_to_markdown_lines(cls, data: Any, scope: str) -> List[str]:
        lines = []
        if isinstance(data, dict):
            # Try to resolve nested content or entities
            content = data.get("content", data)
            entities = data.get("entities", {})
            
            if scope == "Entire Workspace":
                # Aggregate deliverables
                deliverables = data.get("deliverables", {})
                for doc_name, doc_val in deliverables.items():
                    lines.append(f"## {doc_name}")
                    lines.extend(cls._convert_to_markdown_lines(doc_val, doc_name))
                    lines.append("\n---\n")
                return lines
                
            if isinstance(content, dict):
                # Flat sections document
                for k, v in content.items():
                    lines.append(f"### {k}")
                    lines.append(str(v))
                    lines.append("")
            elif isinstance(content, list):
                # List of sections/items
                for item in content:
                    lines.append(f"- {item}")
            else:
                lines.append(str(content))
                
            # Handle structured entities if present
            if entities:
                if "phases" in entities:
                    lines.append("### Roadmap Phases")
                    for p in entities["phases"]:
                        lines.append(f"#### {p.get('id')} - {p.get('phase')} ({p.get('quarter')})")
                        lines.append(f"**Objectives:** {', '.join(p.get('objectives', []))}")
                        lines.append(f"**Milestones:** {', '.join(p.get('milestones', []))}")
                        lines.append("")
                if "tasks" in entities:
                    lines.append("### Jira Tasks")
                    for t in entities["tasks"]:
                        lines.append(f"- **{t.get('id')}**: {t.get('title')} ({t.get('type')}) - {t.get('priority')}")
                        lines.append(f"  Description: {t.get('description')}")
        elif isinstance(data, list):
            for item in data:
                lines.append(f"- {item}")
        else:
            lines.append(str(data))
            
        return lines

    @classmethod
    def _insert_docx_content(cls, doc: DocxDocument, data: Any, scope: str) -> None:
        title_color = RGBColor(30, 58, 138)
        text_color = RGBColor(51, 65, 85)
        
        if scope == "Entire Workspace":
            deliverables = data.get("deliverables", {})
            for doc_name, doc_val in deliverables.items():
                doc.add_heading(doc_name, level=1).runs[0].font.color.rgb = title_color
                cls._insert_docx_content(doc, doc_val, doc_name)
                doc.add_page_break()
            return
            
        content = data.get("content", data) if isinstance(data, dict) else data
        entities = data.get("entities", {}) if isinstance(data, dict) else {}
        
        if isinstance(content, dict):
            for k, v in content.items():
                p_head = doc.add_heading(k, level=2)
                p_head.runs[0].font.color.rgb = title_color
                doc.add_paragraph(str(v)).runs[0].font.color.rgb = text_color
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style='List Bullet').runs[0].font.color.rgb = text_color
        else:
            doc.add_paragraph(str(content)).runs[0].font.color.rgb = text_color
            
        # Insert entities details
        if entities:
            if "phases" in entities:
                doc.add_heading("Roadmap Phases Details", level=2).runs[0].font.color.rgb = title_color
                for p in entities["phases"]:
                    p_text = f"[{p.get('id')}] {p.get('phase')} ({p.get('quarter')})\nObjectives: {', '.join(p.get('objectives', []))}"
                    doc.add_paragraph(p_text).runs[0].font.color.rgb = text_color
            if "tasks" in entities:
                doc.add_heading("Jira Tasks Breakdown", level=2).runs[0].font.color.rgb = title_color
                for t in entities["tasks"]:
                    p_text = f"- {t.get('id')}: {t.get('title')} [{t.get('type')}] Priority: {t.get('priority')}"
                    doc.add_paragraph(p_text).runs[0].font.color.rgb = text_color

    @classmethod
    def _insert_pdf_content(cls, story: List[Any], data: Any, scope: str, style_heading: ParagraphStyle, style_body: ParagraphStyle) -> None:
        if scope == "Entire Workspace":
            deliverables = data.get("deliverables", {})
            for doc_name, doc_val in deliverables.items():
                story.append(Paragraph(doc_name, style_heading))
                story.append(Spacer(1, 10))
                cls._insert_pdf_content(story, doc_val, doc_name, style_heading, style_body)
                story.append(PageBreak())
            return
            
        content = data.get("content", data) if isinstance(data, dict) else data
        entities = data.get("entities", {}) if isinstance(data, dict) else {}
        
        if isinstance(content, dict):
            for k, v in content.items():
                story.append(Paragraph(k, style_heading))
                story.append(Paragraph(str(v).replace("\n", "<br/>"), style_body))
                story.append(Spacer(1, 10))
        elif isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"• {item}", style_body))
        else:
            story.append(Paragraph(str(content).replace("\n", "<br/>"), style_body))
            
        # Entities detail formatting
        if entities:
            if "phases" in entities:
                story.append(Paragraph("Roadmap Phases", style_heading))
                for p in entities["phases"]:
                    text = f"<b>{p.get('id')} — {p.get('phase')} ({p.get('quarter')})</b><br/>Objectives: {', '.join(p.get('objectives', []))}"
                    story.append(Paragraph(text, style_body))
            if "tasks" in entities:
                story.append(Paragraph("Jira Tasks", style_heading))
                for t in entities["tasks"]:
                    text = f"• <b>{t.get('id')}</b>: {t.get('title')} ({t.get('type')}) — <i>{t.get('priority')}</i>"
                    story.append(Paragraph(text, style_body))
